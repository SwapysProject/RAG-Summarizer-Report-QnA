"""
Utility functions for document processing across multiple formats
"""
import os
import io
from pathlib import Path
from typing import List, Dict, Any, Tuple
import tempfile

# PDF Processing
import PyPDF2
import fitz  # PyMuPDF

# Word Documents
from docx import Document

# Excel
import pandas as pd

# Images
from PIL import Image
import pytesseract

from loguru import logger
import config


class DocumentProcessor:
    """Handle multi-format document processing"""
    
    def __init__(self):
        self.supported_formats = config.SUPPORTED_EXTENSIONS
        
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract text, tables, and images
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted content
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self.process_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.process_word(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            return self.process_excel(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg']:
            return self.process_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def process_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """Extract text, tables, and images from PDF"""
        result = {
            'text': [],
            'tables': [],
            'images': [],
            'metadata': {}
        }
        
        doc = None
        try:
            # Use PyMuPDF for comprehensive extraction
            doc = fitz.open(pdf_path)
            
            # Store page count before closing
            num_pages = len(doc)
            logger.info(f"Processing PDF with {num_pages} pages: {pdf_path}")
            
            for page_num, page in enumerate(doc):
                try:
                    # Extract text
                    text = page.get_text()
                    if text.strip():
                        result['text'].append({
                            'page': page_num + 1,
                            'content': text
                        })
                except Exception as text_err:
                    logger.warning(f"Failed to extract text on page {page_num + 1}: {text_err}")
                
                try:
                    # Extract images
                    image_list = page.get_images()
                    for img_index, img in enumerate(image_list):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # Save image to temp location
                        temp_img_path = os.path.join(
                            config.UPLOADS_DIR,
                            f"extracted_img_p{page_num+1}_{img_index}.{image_ext}"
                        )
                        
                        with open(temp_img_path, "wb") as img_file:
                            img_file.write(image_bytes)
                        
                        result['images'].append({
                            'page': page_num + 1,
                            'path': temp_img_path,
                            'format': image_ext
                        })
                except Exception as img_err:
                    logger.warning(f"Failed to extract images on page {page_num + 1}: {img_err}")
                
                # Extract tables (basic detection)
                try:
                    tables = page.find_tables()
                    if tables:
                        for table_index, table in enumerate(tables):
                            try:
                                df = table.to_pandas()
                                result['tables'].append({
                                    'page': page_num + 1,
                                    'data': df.to_dict('records'),
                                    'dataframe': df
                                })
                            except Exception as table_err:
                                logger.warning(f"Failed to extract table {table_index} on page {page_num + 1}: {table_err}")
                                continue
                except Exception as tables_err:
                    logger.warning(f"Failed to find tables on page {page_num + 1}: {tables_err}")
            
            # Metadata (use stored page count)
            result['metadata'] = {
                'filename': Path(pdf_path).name,
                'num_pages': num_pages,
                'format': 'PDF'
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise
        finally:
            # Ensure document is closed
            if doc is not None:
                doc.close()
        
        return result
    
    def process_word(self, docx_path: str) -> Dict[str, Any]:
        """Extract text, tables, and images from Word document"""
        result = {
            'text': [],
            'tables': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            doc = Document(docx_path)
            
            # Extract paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            result['text'].append({
                'page': 1,
                'content': '\n'.join(full_text)
            })
            
            # Extract tables
            for table_index, table in enumerate(doc.tables):
                data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    data.append(row_data)
                
                # Convert to DataFrame
                if data:
                    df = pd.DataFrame(data[1:], columns=data[0])
                    result['tables'].append({
                        'index': table_index,
                        'data': df.to_dict('records'),
                        'dataframe': df
                    })
            
            # Extract images (from document parts)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    img = rel.target_part.blob
                    temp_img_path = os.path.join(
                        config.UPLOADS_DIR,
                        f"word_img_{len(result['images'])}.png"
                    )
                    with open(temp_img_path, "wb") as img_file:
                        img_file.write(img)
                    
                    result['images'].append({
                        'index': len(result['images']),
                        'path': temp_img_path
                    })
            
            result['metadata'] = {
                'filename': Path(docx_path).name,
                'format': 'DOCX'
            }
            
        except Exception as e:
            logger.error(f"Error processing Word document: {e}")
            raise
        
        return result
    
    def process_excel(self, excel_path: str) -> Dict[str, Any]:
        """Extract tables from Excel file"""
        result = {
            'text': [],
            'tables': [],
            'images': [],
            'metadata': {}
        }
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(excel_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)
                
                # Convert to text representation
                text_repr = f"Sheet: {sheet_name}\n" + df.to_string()
                result['text'].append({
                    'sheet': sheet_name,
                    'content': text_repr
                })
                
                # Store as table
                result['tables'].append({
                    'sheet': sheet_name,
                    'data': df.to_dict('records'),
                    'dataframe': df
                })
            
            result['metadata'] = {
                'filename': Path(excel_path).name,
                'format': 'EXCEL',
                'num_sheets': len(excel_file.sheet_names)
            }
            
        except Exception as e:
            logger.error(f"Error processing Excel: {e}")
            raise
        
        return result
    
    def process_image(self, image_path: str) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        result = {
            'text': [],
            'tables': [],
            'images': [{'path': image_path}],
            'metadata': {}
        }
        
        try:
            # Open image
            img = Image.open(image_path)
            
            # Perform OCR
            text = pytesseract.image_to_string(img)
            
            if text.strip():
                result['text'].append({
                    'content': text
                })
            
            # Convert tuple to string for ChromaDB compatibility
            width, height = img.size
            result['metadata'] = {
                'filename': Path(image_path).name,
                'format': 'IMAGE',
                'width': width,
                'height': height,
                'size': f"{width}x{height}"
            }
            
        except Exception as e:
            logger.warning(f"OCR failed for image: {e}. Skipping text extraction.")
            # Still return image path even if OCR fails
        
        return result
    
    def extract_all_text(self, processed_doc: Dict[str, Any]) -> str:
        """Combine all text from processed document"""
        all_text = []
        
        for text_item in processed_doc.get('text', []):
            all_text.append(text_item.get('content', ''))
        
        return '\n\n'.join(all_text)
